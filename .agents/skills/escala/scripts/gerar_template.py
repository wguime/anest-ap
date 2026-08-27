#!/usr/bin/env python3
"""
Gera o template docx UNIFICADO de escala mensal (sobreaviso materno + hospitais).

Uma tabela, uma linha por dia, um valor por célula — desenhado para ser
trivial de preencher (a coordenadora só digita nomes nas células vazias) e
à prova de ambiguidade na hora de eu parsear (sem célula multi-linha, sem
domingo embutido no sábado como no export antigo do Numbers).

Regras de quais slots de hospital se aplicam (espelham src/data/hospitaisTecnicas2026.js):
  - SOBREAVISO: todo dia.
  - UNIMED  (07–15): sábados e feriados. Domingo não tem.
  - HRO     (07–15): sábados, domingos e feriados.
  - PLANTÃO (15–23): sábados, domingos e feriados.
Células de hospital que NÃO se aplicam no dia vêm pré-preenchidas com "—".
Feriados são pré-rotulados a partir de FERIADO_LABELS (plantao2026.js).

Uso:
  python3 gerar_template.py 2026-08                       # salva na pasta padrão
  python3 gerar_template.py 2026-08 "/caminho/custom.docx"  # salva onde quiser
"""
import sys, re, calendar, os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/Users/guilherme/dev/anest"
PLANTAO_JS = f"{REPO}/src/data/plantao2026.js"
# Pasta padrão dos modelos preenchíveis (onde o usuário busca/anexa os docx).
OUT_DIR = "/Users/guilherme/Documents/IA/Escalas funcinárias"

PT_DIAS = ['SEGUNDA', 'TERÇA', 'QUARTA', 'QUINTA', 'SEXTA', 'SÁBADO', 'DOMINGO']
PT_MESES = ['', 'JANEIRO', 'FEVEREIRO', 'MARÇO', 'ABRIL', 'MAIO', 'JUNHO',
            'JULHO', 'AGOSTO', 'SETEMBRO', 'OUTUBRO', 'NOVEMBRO', 'DEZEMBRO']
FUNCIONARIAS = ['Marta', 'Renata', 'Luciana', 'Elisete', 'Saionara', 'Mari']

COLS = ['DATA', 'DIA', 'SOBREAVISO', 'UNIMED (07-15)', 'HRO (07-15)',
        'PLANTÃO PAGO (15-23)', 'FERIADO']
NA = '—'  # marcador "não se aplica" — o importador ignora


def load_feriado_labels():
    """Extrai FERIADO_LABELS de plantao2026.js (fonte única da verdade)."""
    txt = open(PLANTAO_JS, encoding='utf-8').read()
    i = txt.find('FERIADO_LABELS')
    if i < 0:
        return {}
    block = txt[i: txt.find('}', i)]
    return dict(re.findall(r"'(\d{4}-\d{2}-\d{2})':\s*'([^']+)'", block))


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell(cell, text, bold=False, align='left', gray=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    if gray:
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def main():
    if len(sys.argv) < 2:
        print("uso: gerar_template.py YYYY-MM [output.docx]", file=sys.stderr)
        sys.exit(1)
    ym = sys.argv[1]
    if len(sys.argv) >= 3:
        out = sys.argv[2]
    else:
        os.makedirs(OUT_DIR, exist_ok=True)
        out = os.path.join(OUT_DIR, f"Escala {ym}.docx")
    yy, mm = (int(x) for x in ym.split('-'))
    labels = load_feriado_labels()
    ndays = calendar.monthrange(yy, mm)[1]

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"ESCALA {PT_MESES[mm]} {yy} — SOBREAVISO MATERNO + HOSPITAIS")
    r.bold = True
    r.font.size = Pt(13)

    head = doc.add_paragraph()
    head.add_run("COMO PREENCHER").bold = True
    head.runs[0].font.size = Pt(11)

    linhas = [
        "Preencha apenas os NOMES. As datas, os dias da semana e os feriados já estão prontos.",
        "Coluna SOBREAVISO: preencha TODOS os dias (uma pessoa por dia).",
        "Colunas UNIMED, HRO e PLANTÃO PAGO: só nas linhas verdes (sábados, domingos e feriados).",
        f'Onde aparecer "{NA}", não precisa preencher — pode deixar como está.',
        "Use somente estes nomes: " + ", ".join(FUNCIONARIAS) + ".",
    ]
    for txt in linhas:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("•  " + txt)
        r.font.size = Pt(10)
    doc.add_paragraph()  # respiro antes da tabela

    table = doc.add_table(rows=1, cols=len(COLS))
    table.style = 'Table Grid'
    for c, name in zip(table.rows[0].cells, COLS):
        set_cell(c, name, bold=True, align='center')
        shade(c, 'D9D9D9')

    for d in range(1, ndays + 1):
        dt = date(yy, mm, d)
        key = dt.isoformat()
        wd = dt.weekday()  # 0=seg ... 6=dom
        is_sat = wd == 5
        is_sun = wd == 6
        is_fer = key in labels
        unimed_ok = is_sat or is_fer
        hro_ok = is_sat or is_sun or is_fer
        plantao_ok = hro_ok
        destaque = is_sat or is_sun or is_fer

        row = table.add_row().cells
        set_cell(row[0], dt.strftime('%d/%m/%Y'), align='center')
        set_cell(row[1], PT_DIAS[wd], align='center')
        set_cell(row[2], '', align='center')                       # sobreaviso: sempre preencher
        set_cell(row[3], '' if unimed_ok else NA, align='center', gray=not unimed_ok)
        set_cell(row[4], '' if hro_ok else NA, align='center', gray=not hro_ok)
        set_cell(row[5], '' if plantao_ok else NA, align='center', gray=not plantao_ok)
        set_cell(row[6], labels.get(key, ''), align='center')      # feriado pré-rotulado
        if destaque:
            for c in row:
                shade(c, 'EAF5EA')  # verde claro: linha que precisa de nomes de hospital

    doc.save(out)
    n_fds = sum(1 for d in range(1, ndays + 1)
                if date(yy, mm, d).weekday() >= 5 or date(yy, mm, d).isoformat() in labels)
    print(f"OK: {out}")
    print(f"  {ndays} dias (linhas) | {n_fds} linhas de hospital (FDS+feriados)")
    fers = [f"{k} {v}" for k, v in sorted(labels.items()) if k.startswith(ym)]
    print(f"  feriados no mês: {fers if fers else 'nenhum'}")


if __name__ == '__main__':
    main()
